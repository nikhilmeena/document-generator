"""Fill Word templates by replacing yellow-highlighted placeholders."""

from __future__ import annotations

from pathlib import Path
from shutil import copyfile
from tempfile import NamedTemporaryFile
from typing import Union

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from brs_parser import PurposeSection
from models.project_context import ProjectContext

Replacement = Union[str, list[PurposeSection]]


def _is_yellow_element(run_element) -> bool:
    r_pr = run_element.find(qn("w:rPr"))
    if r_pr is None:
        return False
    highlight = r_pr.find(qn("w:highlight"))
    if highlight is not None:
        value = (highlight.get(qn("w:val")) or "").lower()
        return value == "yellow"
    return False


def _is_yellow_run(run) -> bool:
    if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
        return True
    return _is_yellow_element(run._element)


def _run_text(run_element) -> str:
    return "".join(node.text or "" for node in run_element.iter(qn("w:t")))


def _set_run_text(run_element, text: str) -> None:
    text_nodes = run_element.findall(qn("w:t"))
    if not text_nodes:
        text_element = run_element.makeelement(qn("w:t"))
        text_element.text = text
        run_element.append(text_element)
        return
    text_nodes[0].text = text
    for node in text_nodes[1:]:
        node.text = ""


def _paragraph_non_yellow_text(paragraph: Paragraph) -> str:
    parts: list[str] = []
    for run_element in paragraph._element.iter(qn("w:r")):
        if _is_yellow_element(run_element):
            continue
        text = _run_text(run_element)
        if text:
            parts.append(text)
    return "".join(parts).strip()


def _replace_yellow_in_paragraph(paragraph: Paragraph, replacement: str) -> bool:
    yellow_runs = [
        run_element
        for run_element in paragraph._element.iter(qn("w:r"))
        if _is_yellow_element(run_element)
    ]
    if not yellow_runs:
        return False

    non_yellow_text = _paragraph_non_yellow_text(paragraph)

    if non_yellow_text and non_yellow_text.rstrip().endswith(":"):
        _set_run_text(yellow_runs[0], replacement)
        for run_element in yellow_runs[1:]:
            _set_run_text(run_element, "")
        return True

    if non_yellow_text:
        _set_run_text(yellow_runs[0], replacement)
        for run_element in yellow_runs[1:]:
            _set_run_text(run_element, "")
        return True

    _set_run_text(yellow_runs[0], replacement)
    for run_element in yellow_runs[1:]:
        _set_run_text(run_element, "")
    return True


def _replace_yellow_with_sections(
    paragraph: Paragraph,
    sections: list[PurposeSection],
) -> None:
    """Replace a yellow paragraph with numbered bullets, bold titles, and line spacing."""
    if not sections:
        _replace_yellow_in_paragraph(paragraph, "")
        return

    document = paragraph.part.document
    parent = paragraph._element.getparent()
    insert_idx = list(parent).index(paragraph._element)
    parent.remove(paragraph._element)

    created_elements = []
    for section_index, section in enumerate(sections):
        if section_index > 0:
            created_elements.append(document.add_paragraph()._element)

        header = document.add_paragraph()
        header.add_run(f"{section.number}. ")
        title_run = header.add_run(section.title)
        title_run.bold = True
        created_elements.append(header._element)

        if section.lines:
            created_elements.append(document.add_paragraph()._element)
            for line in section.lines:
                created_elements.append(document.add_paragraph(line)._element)

    body = document.element.body
    for offset, element in enumerate(created_elements):
        if element.getparent() is not None:
            element.getparent().remove(element)
        parent.insert(insert_idx + offset, element)


def _apply_replacement(paragraph: Paragraph, replacement: Replacement) -> None:
    if isinstance(replacement, list):
        _replace_yellow_with_sections(paragraph, replacement)
    else:
        _replace_yellow_in_paragraph(paragraph, replacement)


def _iter_cell_paragraphs(cell: _Cell):
    seen: set[int] = set()
    for paragraph in cell.paragraphs:
        key = id(paragraph._element)
        if key in seen:
            continue
        seen.add(key)
        yield paragraph


def _iter_body_paragraphs(document: Document):
    """Walk paragraphs in document body order (matches template yellow sequence)."""
    for child in document.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_cell_paragraphs(cell)


def _apply_replacements(
    template_path: Path,
    output_path: Path,
    replacements: list[Replacement],
) -> None:
    copyfile(template_path, output_path)
    document = Document(output_path)
    index = 0

    for paragraph in _iter_body_paragraphs(document):
        if not any(
            _is_yellow_element(run_element)
            for run_element in paragraph._element.iter(qn("w:r"))
        ):
            continue
        if index >= len(replacements):
            break
        _apply_replacement(paragraph, replacements[index])
        index += 1

    document.save(output_path)


def _approach_replacements(context: ProjectContext) -> list[Replacement]:
    return [
        context.cr_number,
        context.deployment_date,
        context.spoc_name,
        context.deployment_date,
        context.spoc_name,
        context.deployment_date,
        context.spoc_name,
        context.deployment_date,
        context.cr_number,
        context.introduction_sections(),
        context.all_requirement_sections(),
    ]


def _impact_replacements(context: ProjectContext) -> list[Replacement]:
    summary = context.introduction_sections()
    return [
        context.cr_number,
        summary,
        summary,
    ]


def _uat_replacements(context: ProjectContext) -> list[Replacement]:
    body_sections = context.uat_body_sections()
    while len(body_sections) < 3:
        body_sections.append([])

    return [
        context.cr_number,
        context.deployment_title(),
        body_sections[0] if body_sections[0] else "",
        body_sections[1] if body_sections[1] else "",
        body_sections[2] if body_sections[2] else "",
    ]


def generate_approach_document(template_path: Path, context: ProjectContext) -> Path:
    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        output_path = Path(temp_file.name)
    _apply_replacements(template_path, output_path, _approach_replacements(context))
    return output_path


def generate_impact_document(template_path: Path, context: ProjectContext) -> Path:
    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        output_path = Path(temp_file.name)
    _apply_replacements(template_path, output_path, _impact_replacements(context))
    return output_path


def generate_uat_deployment_note(template_path: Path, context: ProjectContext) -> Path:
    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        output_path = Path(temp_file.name)
    _apply_replacements(template_path, output_path, _uat_replacements(context))
    return output_path
